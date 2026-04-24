"""Word tool — Create professional .docx files via python-docx.

Produces executive-quality documents with inline formatting, tables,
headers/footers, page numbers, and smart professional defaults.
"""

import os
import re
from datetime import date
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from neo.tools.paths import resolve_path, _validate_write_path

# ---------------------------------------------------------------------------
# Defaults
# ---------------------------------------------------------------------------
_DEFAULT_FONT = "Calibri"
_DEFAULT_FONT_SIZE = 11
_DEFAULT_HEADING_FONT = "Calibri"
_DEFAULT_LINE_SPACING = 1.15
_DEFAULT_MARGINS = {"top": 1, "bottom": 1, "left": 1, "right": 1}

# Table styling
_TABLE_HEADER_BG = "2B5797"
_TABLE_HEADER_FONT_COLOR = "FFFFFF"
_TABLE_SHADE_BG = "EDF2F9"


# ---------------------------------------------------------------------------
# Inline formatting parser
# ---------------------------------------------------------------------------
def _parse_inline_formatting(paragraph, text: str) -> None:
    """Parse markdown-like inline formatting and add runs to a paragraph.

    Supports: **bold**, *italic*, __underline__, `code`,
    and combinations like ***bold italic***.
    """
    # Pattern matches: **bold**, *italic*, __underline__, `code`, or plain text
    pattern = re.compile(
        r"(\*\*\*(.+?)\*\*\*)"  # ***bold italic***
        r"|(\*\*(.+?)\*\*)"  # **bold**
        r"|(\*(.+?)\*)"  # *italic*
        r"|(__(.+?)__)"  # __underline__
        r"|(`(.+?)`)"  # `code`
        r"|([^*_`]+)"  # plain text
    )

    for match in pattern.finditer(text):
        if match.group(2):  # ***bold italic***
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(4):  # **bold**
            run = paragraph.add_run(match.group(4))
            run.bold = True
        elif match.group(6):  # *italic*
            run = paragraph.add_run(match.group(6))
            run.italic = True
        elif match.group(8):  # __underline__
            run = paragraph.add_run(match.group(8))
            run.underline = True
        elif match.group(10):  # `code`
            run = paragraph.add_run(match.group(10))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        elif match.group(11):  # plain text
            paragraph.add_run(match.group(11))


def _set_paragraph_spacing(paragraph, before: int = 0, after: int = 6) -> None:
    """Set paragraph spacing in points."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


# ---------------------------------------------------------------------------
# Table builder
# ---------------------------------------------------------------------------
def _add_table(doc: Document, table_data: dict, style: str = "professional") -> None:
    """Add a professionally styled table to the document.

    table_data: {headers: [...], rows: [[...], ...], style: "professional"}
    """
    headers = table_data.get("headers", [])
    rows = table_data.get("rows", [])
    tbl_style = table_data.get("style", style)

    if not headers:
        return

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(header))
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = _DEFAULT_FONT
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Header cell shading
        if tbl_style in ("professional", "grid"):
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), _TABLE_HEADER_BG)
            shading.set(qn("w:val"), "clear")
            cell._element.get_or_add_tcPr().append(shading)
            run.font.color.rgb = RGBColor.from_string(_TABLE_HEADER_FONT_COLOR)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            if col_idx < len(headers):
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = str(value) if value is not None else ""
                p = cell.paragraphs[0]
                if p.runs:
                    p.runs[0].font.size = Pt(10)
                    p.runs[0].font.name = _DEFAULT_FONT

                # Alternating row shading
                if tbl_style == "professional" and row_idx % 2 == 1:
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), _TABLE_SHADE_BG)
                    shading.set(qn("w:val"), "clear")
                    cell._element.get_or_add_tcPr().append(shading)

    # Set column widths evenly
    total_width = Inches(6.5)  # Standard page width minus margins
    col_width = total_width // len(headers)
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_width

    doc.add_paragraph("")  # Spacing after table


# ---------------------------------------------------------------------------
# Headers, footers, page numbers
# ---------------------------------------------------------------------------
def _add_page_numbers(doc: Document) -> None:
    """Add page numbers to the footer (centered)."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style = doc.styles["Normal"]

        # PAGE field
        run = p.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        run._element.append(fldChar1)

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = " PAGE "
        run._element.append(instrText)

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._element.append(fldChar2)


def _add_headers_footers(doc: Document, config: dict) -> None:
    """Add custom headers and footers to the document.

    Supports placeholders: {date}, {page}, {pages}
    """
    today = date.today().strftime("%Y-%m-%d")

    for section in doc.sections:
        # Header
        header_left = config.get("header_left", "")
        header_right = config.get("header_right", "")

        if header_left or header_right:
            header = section.header
            header.is_linked_to_previous = False
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.clear()

            if header_left:
                run = p.add_run(header_left.replace("{date}", today))
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

            if header_right:
                # Add tab stop for right alignment
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if header_left:
                    run = p.add_run("\t\t")  # Push to right
                run = p.add_run(header_right.replace("{date}", today))
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        # Footer
        footer_center = config.get("footer_center", "")
        footer_right = config.get("footer_right", "")

        if footer_center or footer_right:
            footer = section.footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.clear()

            if footer_center:
                text = footer_center.replace("{date}", today)
                # Handle {page} placeholder with field code
                if "{page}" in text:
                    parts = text.split("{page}")
                    run = p.add_run(parts[0])
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

                    # PAGE field
                    run2 = p.add_run()
                    fld1 = OxmlElement("w:fldChar")
                    fld1.set(qn("w:fldCharType"), "begin")
                    run2._element.append(fld1)
                    instr = OxmlElement("w:instrText")
                    instr.set(qn("xml:space"), "preserve")
                    instr.text = " PAGE "
                    run2._element.append(instr)
                    fld2 = OxmlElement("w:fldChar")
                    fld2.set(qn("w:fldCharType"), "end")
                    run2._element.append(fld2)

                    if len(parts) > 1:
                        remaining = parts[1]
                        if "{pages}" in remaining:
                            sub_parts = remaining.split("{pages}")
                            run3 = p.add_run(sub_parts[0])
                            run3.font.size = Pt(9)
                            run3.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                            # NUMPAGES field
                            run4 = p.add_run()
                            fld3 = OxmlElement("w:fldChar")
                            fld3.set(qn("w:fldCharType"), "begin")
                            run4._element.append(fld3)
                            instr2 = OxmlElement("w:instrText")
                            instr2.set(qn("xml:space"), "preserve")
                            instr2.text = " NUMPAGES "
                            run4._element.append(instr2)
                            fld4 = OxmlElement("w:fldChar")
                            fld4.set(qn("w:fldCharType"), "end")
                            run4._element.append(fld4)
                            if len(sub_parts) > 1:
                                run5 = p.add_run(sub_parts[1])
                                run5.font.size = Pt(9)
                                run5.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                        else:
                            run3 = p.add_run(remaining)
                            run3.font.size = Pt(9)
                            run3.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                else:
                    run = p.add_run(text)
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

                p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_toc(doc: Document) -> None:
    """Add a Table of Contents field to the document.

    The TOC will be populated when the document is opened in Word
    and the user updates fields (Ctrl+A, F9).
    """
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["Normal"]

    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._element.append(fldChar1)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._element.append(instrText)

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run._element.append(fldChar2)

    # Placeholder text until fields are updated
    run2 = paragraph.add_run("[Update this Table of Contents: select all with Ctrl+A, then press F9]")
    run2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run2.font.italic = True

    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run2._element.append(fldChar3)

    doc.add_paragraph("")  # Spacing after TOC


# ---------------------------------------------------------------------------
# Markdown table detection
# ---------------------------------------------------------------------------
def _detect_markdown_table(lines: list[str], start_idx: int) -> tuple[dict | None, int]:
    """Detect a markdown table starting at the given line index.

    Returns (table_data, lines_consumed) or (None, 0).
    """
    if start_idx >= len(lines):
        return None, 0

    line = lines[start_idx].strip()
    if not line.startswith("|") or not line.endswith("|"):
        return None, 0

    # Parse header row
    headers = [cell.strip() for cell in line.split("|")[1:-1]]
    if not headers:
        return None, 0

    consumed = 1

    # Check for separator row (|---|---|)
    if start_idx + 1 < len(lines):
        sep_line = lines[start_idx + 1].strip()
        if sep_line.startswith("|") and all(
            c in "-|: " for c in sep_line
        ):
            consumed = 2
        else:
            return None, 0
    else:
        return None, 0

    # Parse data rows
    rows = []
    idx = start_idx + consumed
    while idx < len(lines):
        data_line = lines[idx].strip()
        if not data_line.startswith("|") or not data_line.endswith("|"):
            break
        row = [cell.strip() for cell in data_line.split("|")[1:-1]]
        rows.append(row)
        consumed += 1
        idx += 1

    if rows:
        return {"headers": headers, "rows": rows}, consumed

    return None, 0


# ---------------------------------------------------------------------------
# Document styling
# ---------------------------------------------------------------------------
def _apply_document_formatting(doc: Document, formatting: dict | None) -> None:
    """Apply document-level formatting defaults."""
    fmt = formatting or {}

    font_name = fmt.get("font", _DEFAULT_FONT)
    font_size = fmt.get("font_size", _DEFAULT_FONT_SIZE)
    line_spacing = fmt.get("line_spacing", _DEFAULT_LINE_SPACING)
    margins = fmt.get("margins", _DEFAULT_MARGINS)

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style.paragraph_format.line_spacing = line_spacing
    style.paragraph_format.space_after = Pt(6)

    # Set heading fonts
    heading_font = fmt.get("heading_font", _DEFAULT_HEADING_FONT)
    for level in range(1, 4):
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            h_style = doc.styles[style_name]
            h_style.font.name = heading_font
            h_style.font.color.rgb = RGBColor(0x2B, 0x57, 0x97)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(margins.get("top", 1))
        section.bottom_margin = Inches(margins.get("bottom", 1))
        section.left_margin = Inches(margins.get("left", 1))
        section.right_margin = Inches(margins.get("right", 1))


# ---------------------------------------------------------------------------
# Content string parser (enhanced, backward compatible)
# ---------------------------------------------------------------------------
def _parse_content_string(doc: Document, content: str) -> int:
    """Parse a content string with enhanced markdown support.

    Returns the number of headings found (for auto-TOC decision).
    """
    lines = content.split("\n")
    heading_count = 0
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        # Page break
        if line.strip() == "---":
            doc.add_page_break()
            i += 1
            continue

        # Markdown table detection
        table_data, consumed = _detect_markdown_table(lines, i)
        if table_data:
            _add_table(doc, table_data)
            i += consumed
            continue

        # Empty line
        if not line:
            doc.add_paragraph("")
            i += 1
            continue

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
            heading_count += 1
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
            heading_count += 1
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
            heading_count += 1

        # Bullet list
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _parse_inline_formatting(p, line[2:])

        # Numbered list
        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line)
            p = doc.add_paragraph(style="List Number")
            _parse_inline_formatting(p, text)

        # Normal paragraph with inline formatting
        else:
            p = doc.add_paragraph()
            _parse_inline_formatting(p, line)

        i += 1

    return heading_count


# ---------------------------------------------------------------------------
# Sections-based builder
# ---------------------------------------------------------------------------
def _build_from_sections(doc: Document, sections: list[dict]) -> int:
    """Build document from structured sections.

    Returns the number of headings found.
    """
    heading_count = 0

    for section in sections:
        # Page break
        if section.get("page_break"):
            doc.add_page_break()
            continue

        # Heading
        heading = section.get("heading", "")
        level = section.get("level", 1)
        if heading:
            doc.add_heading(heading, level=level)
            heading_count += 1

        # Body text with inline formatting
        body = section.get("body", "")
        if body:
            for para_text in body.split("\n"):
                if para_text.strip():
                    p = doc.add_paragraph()
                    _parse_inline_formatting(p, para_text.strip())

        # Bullet list
        bullets = section.get("bullets", [])
        for bullet in bullets:
            p = doc.add_paragraph(style="List Bullet")
            _parse_inline_formatting(p, bullet)

        # Numbered list
        numbered = section.get("numbered_list", [])
        for item in numbered:
            p = doc.add_paragraph(style="List Number")
            _parse_inline_formatting(p, item)

        # Table
        table_data = section.get("table")
        if table_data:
            _add_table(doc, table_data, style=table_data.get("style", "professional"))

    return heading_count


# ---------------------------------------------------------------------------
# Main function
# ---------------------------------------------------------------------------
def create_document(
    title: str,
    content: str = "",
    output_path: str | None = None,
    sections: list[dict] | None = None,
    formatting: dict | None = None,
    headers_footers: dict | None = None,
) -> str:
    """Create a professional Word document.

    Args:
        title: Filename (without extension) or full path.
        content: Document body with enhanced markdown support.
            - # Heading 1, ## Heading 2, ### Heading 3
            - - bullet items
            - 1. numbered items
            - **bold**, *italic*, __underline__, `code`
            - --- for page breaks
            - |H1|H2| / |---|---| / |d1|d2| for tables
        output_path: Directory where the file should be saved.
        sections: Structured alternative to content string. List of section dicts:
            - heading (str), level (int), body (str), bullets (list),
              numbered_list (list), table ({headers, rows, style}),
              page_break (bool)
        formatting: Document-level styles:
            - font, font_size, heading_font, line_spacing,
              margins ({top, bottom, left, right} in inches),
              page_numbers (bool), toc (bool)
        headers_footers: Header/footer content:
            - header_left, header_right, footer_center, footer_right
            - Supports {date}, {page}, {pages} placeholders

    Returns:
        Absolute path to the created .docx file.
    """
    doc = Document()

    # Apply formatting defaults
    _apply_document_formatting(doc, formatting)

    # Add custom headers/footers
    if headers_footers:
        _add_headers_footers(doc, headers_footers)

    # Document title
    doc.add_heading(os.path.basename(title), level=0)

    fmt = formatting or {}
    heading_count = 0

    # Table of Contents (if requested or auto for 4+ headings)
    add_toc = fmt.get("toc", False)

    if sections:
        # If TOC requested, add before sections
        if add_toc:
            _add_toc(doc)

        heading_count = _build_from_sections(doc, sections)

        # Auto-add TOC if many headings and not already added
        if not add_toc and heading_count >= 4:
            # Insert TOC after the title (position 1)
            # Since we can't easily insert at position, we note this for the user
            pass
    elif content:
        if add_toc:
            _add_toc(doc)

        heading_count = _parse_content_string(doc, content)

    # Page numbers (default True for documents with multiple sections)
    add_page_numbers = fmt.get("page_numbers", heading_count >= 2)
    if add_page_numbers:
        _add_page_numbers(doc)

    # Save
    file_path = resolve_path(title, ".docx", output_dir=output_path)
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    doc.save(file_path)
    return file_path


# ---------------------------------------------------------------------------
# Collaborative editing
# ---------------------------------------------------------------------------
def edit_document(
    file_path: str,
    operations: list[dict] | None = None,
    save_as: str | None = None,
) -> str:
    """Edit an existing Word document — replace text, add/delete paragraphs, tables.

    Args:
        file_path: Path to the existing .docx file.
        operations: List of editing operations. Each dict has a "type" key:
            - replace_text: {type, find, replace} — find/replace across all paragraphs
            - add_paragraph: {type, text} — append a paragraph (supports inline formatting)
            - add_heading: {type, text, level} — append a heading (level 1-3)
            - add_bullet: {type, text} — append a bullet point
            - add_table: {type, headers, rows} — append a table
            - delete_paragraph: {type, index} — delete paragraph by index (0-based)
            - replace_paragraph: {type, index, text} — replace paragraph text by index
        save_as: Optional alternate save path. If omitted, overwrites original.

    Returns:
        Absolute path to the saved file with a summary of changes.
    """
    from neo.tools.file_reader import _validate_read_path

    real_path = _validate_read_path(file_path)
    doc = Document(real_path)

    changes: list[str] = []

    if operations:
        # Collect delete indices first to process in reverse later
        deletes: list[int] = []

        for op in operations:
            op_type = op.get("type", "")

            if op_type == "replace_text":
                find_str = op.get("find", "")
                replace_str = op.get("replace", "")
                if find_str:
                    count = 0
                    for para in doc.paragraphs:
                        if find_str in para.text:
                            # Preserve formatting by replacing in runs
                            for run in para.runs:
                                if find_str in run.text:
                                    run.text = run.text.replace(find_str, replace_str)
                                    count += 1
                    changes.append(f"Replaced '{find_str}' → '{replace_str}' ({count} runs)")

            elif op_type == "add_paragraph":
                text = op.get("text", "")
                p = doc.add_paragraph()
                _parse_inline_formatting(p, text)
                changes.append("Added paragraph")

            elif op_type == "add_heading":
                text = op.get("text", "")
                level = op.get("level", 1)
                doc.add_heading(text, level=min(max(level, 0), 9))
                changes.append(f"Added heading (level {level})")

            elif op_type == "add_bullet":
                text = op.get("text", "")
                p = doc.add_paragraph(style="List Bullet")
                _parse_inline_formatting(p, text)
                changes.append("Added bullet")

            elif op_type == "add_table":
                table_data = {"headers": op.get("headers", []), "rows": op.get("rows", [])}
                _add_table(doc, table_data)
                changes.append("Added table")

            elif op_type == "delete_paragraph":
                idx = op.get("index", -1)
                if 0 <= idx < len(doc.paragraphs):
                    deletes.append(idx)

            elif op_type == "replace_paragraph":
                idx = op.get("index", -1)
                text = op.get("text", "")
                if 0 <= idx < len(doc.paragraphs):
                    para = doc.paragraphs[idx]
                    # Clear existing runs
                    for run in para.runs:
                        run.text = ""
                    # Clear the paragraph element of child runs
                    for child in list(para._element):
                        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                        if tag == "r":
                            para._element.remove(child)
                    _parse_inline_formatting(para, text)
                    changes.append(f"Replaced paragraph {idx}")

        # Process deletes in reverse order to keep indices stable
        for idx in sorted(deletes, reverse=True):
            if 0 <= idx < len(doc.paragraphs):
                para = doc.paragraphs[idx]
                parent = para._element.getparent()
                parent.remove(para._element)
                changes.append(f"Deleted paragraph {idx}")

    # --- Save ---
    if save_as:
        out_path = os.path.expanduser(save_as)
        if not out_path.endswith(".docx"):
            out_path += ".docx"
        _validate_write_path(out_path)
    else:
        out_path = real_path

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)

    summary = "; ".join(changes) if changes else "No changes applied"
    return f"Saved {out_path} ({summary})"
